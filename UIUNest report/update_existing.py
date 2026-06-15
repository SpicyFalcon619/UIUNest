from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
import os

def update_doc(filename):
    print("Loading doc...")
    doc = Document(filename)

    print("Replacing em dashes...")
    count = 0
    for p in doc.paragraphs:
        for run in p.runs:
            if '—' in run.text:
                run.text = run.text.replace('—', '-')
                count += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if '—' in run.text:
                            run.text = run.text.replace('—', '-')
                            count += 1
    print(f"Replaced {count} em dashes.")

    print("Adding header...")
    for section in doc.sections:
        header = section.header
        if not header.paragraphs:
            p = header.add_paragraph()
        else:
            p = header.paragraphs[0]
            p.clear()
        
        run1 = p.add_run("UIUNest")
        run1.font.bold = True
        
        # In Word, headers typically have a center and right tab stop by default.
        # So "\t\t" pushes the next text to the right margin.
        p.add_run("\t\tPage ")
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        p.runs[-1]._r.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = " PAGE "
        p.runs[-1]._r.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        p.runs[-1]._r.append(fldChar2)

    print("Replacing ERD image...")
    # Find the paragraph with the ERD caption
    for i, p in enumerate(doc.paragraphs):
        if "UIUNest Entity-Relationship Diagram (19 tables)" in p.text:
            if i > 0:
                img_p = doc.paragraphs[i-1]
                # clear the run(s) containing the image
                img_p.clear()
                run = img_p.add_run()
                if os.path.exists("new_erd.png"):
                    run.add_picture("new_erd.png", width=Inches(6.2))
                    print("ERD image replaced.")
                else:
                    print("new_erd.png not found!")
            break

    # Save to the same file
    doc.save(filename)
    print("Done!")

if __name__ == "__main__":
    update_doc("UIUNest_Report_Final.docx")
