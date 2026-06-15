"""
UIUNest Final Report Generator
Generates UIUNest_Report_Final.docx with all sections, screenshots, SQL queries, and current schema.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPORT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_PATH = os.path.join(REPORT_DIR, "UIUNest_Report_Final.docx")

# Colors
NAVY = RGBColor(0x1A, 0x5C, 0x45)   # Emerald green (primary brand color)
GOLD  = RGBColor(0x15, 0x80, 0x3D)   # Darker emerald accent
DARK  = RGBColor(0x1E, 0x1E, 0x2E)   # Near-black
GRAY  = RGBColor(0x64, 0x74, 0x8B)   # Muted gray
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT = RGBColor(0xEE, 0xF7, 0xF2)   # Light emerald background

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    style_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3'}
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = color or NAVY
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(6)
    return p

def add_code_block(doc, code_text, title=None):
    if title:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.font.bold = True
        r.font.color.rgb = NAVY
        r.font.size = Pt(10)
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x16, 0x38, 0x2B)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    # Light background shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'EEF7F2')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    return p

def add_screenshot(doc, filename, caption, width_inches=6.0):
    path = os.path.join(REPORT_DIR, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_inches))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(f"Figure: {caption}")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.italic = True
        cap.runs[0].font.color.rgb = GRAY
        cap.paragraph_format.space_after = Pt(12)
    else:
        add_body(doc, f"[Screenshot: {caption} - file not found: {filename}]", italic=True)

def add_table_header_row(table, headers, bg='1A5C45'):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        set_cell_bg(cell, bg)
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_simple_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    add_table_header_row(table, headers)
    for i, row_data in enumerate(rows):
        row = table.rows[i+1]
        bg = 'FFFFFF' if i % 2 == 0 else 'EEF7F2'
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_text)
            set_cell_bg(cell, bg)
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(9)
    return table

def add_word_toc(doc):
    """Insert a Word TOC field. Opens in Word as a clickable, auto-updating table of contents."""
    # Add the TOC heading
    p_title = doc.add_paragraph()
    r_title = p_title.add_run('Table of Contents')
    r_title.font.size = Pt(18)
    r_title.font.bold = True
    r_title.font.color.rgb = NAVY
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(12)

    # Build the TOC field paragraph
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    # Begin field
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)

    # Field instruction text: TOC with hyperlinks (\h), levels 1-3 (\o), use styles (\u)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instrText)

    # Separate
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar_sep)

    # Placeholder text (shown before Word updates the field)
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_end)

    # Instruction note
    note = doc.add_paragraph()
    nr = note.add_run('[ Right-click this area in Word and select "Update Field" to generate the clickable TOC ]')
    nr.font.size = Pt(9)
    nr.font.italic = True
    nr.font.color.rgb = GRAY
    note.paragraph_format.space_after = Pt(4)


# ==================== BUILD DOCUMENT ====================
doc = Document()

# Page margins
from docx.oxml import OxmlElement
sections = doc.sections
for section in sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

# ======================== COVER PAGE ========================
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover.add_run("\n\n\n")

logo_p = doc.add_paragraph()
logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = logo_p.add_run("UIU")
r.font.size = Pt(48)
r.font.bold = True
r.font.color.rgb = NAVY
r2 = logo_p.add_run("Nest")
r2.font.size = Pt(48)
r2.font.bold = True
r2.font.color.rgb = GOLD

tagline = doc.add_paragraph("Smart Housing & Marketplace Portal for UIU Students")
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tagline.runs[0].font.size = Pt(14)
tagline.runs[0].font.color.rgb = GRAY

doc.add_paragraph("\n")
div = doc.add_paragraph("─" * 60)
div.alignment = WD_ALIGN_PARAGRAPH.CENTER
div.runs[0].font.color.rgb = NAVY

doc.add_paragraph()

info_lines = [
    ("University:", "United International University (UIU)"),
    ("Department:", "Computer Science and Engineering"),
    ("Course:", "Database Management Systems Laboratory"),
    ("Course Code:", "CSE 3522 / CSI 222"),
    ("Report Type:", "Final Project Report"),
    ("Submission:", "June 2026"),
]
for label, value in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{label}  ")
    r.font.bold = True
    r.font.color.rgb = NAVY
    r.font.size = Pt(11)
    r2 = p.add_run(value)
    r2.font.color.rgb = DARK
    r2.font.size = Pt(11)

# ── Team Members ──
doc.add_paragraph()
div2 = doc.add_paragraph()
div2.alignment = WD_ALIGN_PARAGRAPH.CENTER
div2r = div2.add_run("Team Members")
div2r.font.size = Pt(13)
div2r.font.bold = True
div2r.font.color.rgb = NAVY

tm_table = doc.add_table(rows=6, cols=3)
tm_table.style = 'Table Grid'
tm_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
hdr_cells = tm_table.rows[0].cells
for cell, txt in zip(hdr_cells, ['No.', 'Full Name', 'Student ID']):
    set_cell_bg(cell, '1A5C45')
    run = cell.paragraphs[0].add_run(txt)
    run.font.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 5 empty member rows
for i in range(1, 6):
    row = tm_table.rows[i]
    bg = 'FFFFFF' if i % 2 == 1 else 'EEF7F2'
    # No. cell
    set_cell_bg(row.cells[0], bg)
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_no = row.cells[0].paragraphs[0].add_run(str(i))
    r_no.font.size = Pt(10)
    r_no.font.color.rgb = GRAY
    # Name cell
    set_cell_bg(row.cells[1], bg)
    r_name = row.cells[1].paragraphs[0].add_run('')
    r_name.font.size = Pt(10)
    # ID cell
    set_cell_bg(row.cells[2], bg)
    row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_id = row.cells[2].paragraphs[0].add_run('')
    r_id.font.size = Pt(10)

doc.add_page_break()

# ======================== TABLE OF CONTENTS (clickable Word field) ========================
add_word_toc(doc)
doc.add_page_break()

# ======================== SECTION 1: INTRODUCTION ========================
add_heading(doc, "1. Introduction / Overview", level=1)
add_body(doc, """UIUNest is a fully functional, database-driven web application designed exclusively for the students and faculty of United International University (UIU) and the surrounding campus community. The platform serves as a structured ecosystem connecting three types of users - Students, Landlords, and Administrators - through an integrated housing and marketplace system.

The application addresses a core problem that every UIU student living away from home has personally encountered: finding a room or mess near campus currently requires physically walking through surrounding areas, relying on word-of-mouth, or posting on informal Facebook groups where information is unverified, unstructured, and unreliable.

UIUNest replaces this fragmented process with a single platform built around five major components:

  • Property listings with full cost transparency and structured cost breakdowns
  • Flatmate compatibility matching based on user preference profiles
  • A "Seeking" board where students post housing requirements
  • A multi-dimensional review and complaint system for trust and accountability
  • UIUNest Exchange - an integrated student-to-student marketplace for second-hand household items

The project demonstrates advanced relational database concepts including multi-table JOINs, aggregate functions, subqueries, transactional INSERT/UPDATE/DELETE operations, and enforced referential integrity via foreign key constraints - making it an ideal vehicle for a Database Management Systems laboratory course.""")

doc.add_page_break()

# ======================== SECTION 2: MOTIVATION ========================
add_heading(doc, "2. Motivation", level=1)
add_body(doc, "The motivation for UIUNest is rooted in problems the team has experienced firsthand as UIU students. Four distinct real-world problems were identified:")

problems = [
    ("The Hidden Costs Problem", "A landlord near UIU may advertise a room for 3,500 BDT per month. What they do not disclose is the electricity bill on a shared sub-meter, gas bill, building maintenance, caretaker fee, and three months' advance deposit. UIUNest solves this by requiring landlords to complete a structured cost breakdown form. Every charge is stored as a separate database field and displayed transparently, with the total true monthly cost computed automatically."),
    ("The Compatibility Problem", "Shared messes require social compatibility beyond rent. A student who studies until 2am cannot comfortably share a room with someone who sleeps at 10pm. Dietary preferences, guest policies, and cleanliness standards are never discussed through a Facebook post. UIUNest introduces a housing preference profile for every user and computes a compatibility score by comparing preference tables via SQL JOINs."),
    ("The Trust Problem", "Student housing in Bangladesh has a well-known trust deficit. Landlords misrepresent properties, charge undisclosed fees, or refuse to return security deposits. UIUNest addresses this through a multi-dimensional review system, a formal complaint mechanism with admin oversight, and a landlord verification badge requiring document submission."),
    ("The Community Goods Problem", "UIU students accumulate household items - fans, study tables, chairs, mattresses, mini-fridges. UIUNest Exchange is an integrated student-to-student marketplace. Because all users are already in a campus housing context, buyers and sellers are geographically close by default, and a structured offer-and-counteroffer system removes the awkwardness of open price negotiation."),
]
for title, text in problems:
    add_heading(doc, title, level=2)
    add_body(doc, text)

doc.add_page_break()

# ======================== SECTION 3: SIMILAR PROJECTS ========================
add_heading(doc, "3. Similar Projects", level=1)
add_body(doc, "Several existing platforms share partial features with UIUNest, but none were designed for student housing at UIU or in Bangladesh, and none integrate housing search, cost transparency, flatmate compatibility, and a linked marketplace.")

platforms = [
    ("CareTutors (caretutors.com)", "Primary architectural inspiration - a Bangladesh-based two-sided marketplace where one party creates a structured profile and the other searches and filters to find a match. UIUNest mirrors this approach and extends it with a marketplace layer."),
    ("Bproperty (bproperty.com)", "Bangladesh's largest real estate platform. Target market is professionals and families. Does not support mess listings, shared rooms, student-specific cost transparency, or flatmate matching."),
    ("Bikroy.com", "Bangladesh's largest classifieds platform. Rental and second-hand goods sections function as raw, unstructured bulletin boards with no verification, standardized pricing, or compatibility features."),
    ("Facebook Groups", "Large portion of UIU student housing activity. No database structure - posts are unformatted, unverifiable, and quickly buried. UIUNest is the structured, persistent, searchable replacement."),
    ("SpareRoom (UK)", "Technical reference for compatibility matching and review system. Operates in a different market and does not address cost transparency or Bangladesh-specific needs."),
]
for title, text in platforms:
    add_heading(doc, title, level=2)
    add_body(doc, text)

doc.add_page_break()

# ======================== SECTION 4: BENCHMARK ========================
add_heading(doc, "4. Benchmark Analysis", level=1)
add_body(doc, "● = Full support   ◑ = Partial support   ○ = Absent", italic=True)
doc.add_paragraph()

bench_headers = ["Feature", "UIUNest", "Bproperty", "Bikroy", "Facebook", "SpareRoom"]
bench_rows = [
    ["Student-focused housing", "●", "○", "◑", "◑", "◑"],
    ["Mess / shared room support", "●", "○", "○", "◑", "●"],
    ["Itemized cost transparency", "●", "○", "○", "○", "○"],
    ["Flatmate compatibility score", "●", "○", "○", "○", "◑"],
    ["Verified landlord badge", "●", "◑", "○", "○", "◑"],
    ["Multi-dimensional reviews", "●", "◑", "○", "○", "◑"],
    ["Seeking / looking-for posts", "●", "○", "○", "●", "●"],
    ["Live interactive map", "●", "●", "○", "○", "◑"],
    ["Rent history tracker", "●", "○", "○", "○", "○"],
    ["Complaint mechanism", "●", "○", "○", "○", "○"],
    ["Mess bill management module", "●", "○", "○", "○", "○"],
    ["Student goods marketplace", "●", "○", "●", "●", "○"],
    ["Zone-filtered item listings", "●", "○", "○", "○", "○"],
    ["Offer & negotiation system", "●", "○", "◑", "○", "○"],
    ["Housing-linked item listing", "●", "○", "○", "○", "○"],
    ["Admin moderation panel", "●", "◑", "○", "○", "◑"],
    ["UIU campus context", "●", "○", "○", "○", "○"],
    ["In-app notifications", "●", "○", "○", "○", "◑"],
    ["Watchlist / save listings", "●", "●", "○", "○", "●"],
    ["Application tracking system", "●", "○", "○", "○", "◑"],
]
add_simple_table(doc, bench_headers, bench_rows)
doc.add_page_break()

# ======================== SECTION 5: FEATURES ========================
add_heading(doc, "5. Complete Feature List", level=1)
add_body(doc, "UIUNest is organized into the following major feature groups, all of which are fully implemented in the current version:")

features = [
    ("5.1 User Account Management", [
        "Registration and login for three roles: Student, Landlord, Admin",
        "Dual-purpose student role - same account can search for housing AND list a spare seat (peer listing)",
        "Profile: name, email, UIU university ID, gender, phone, profile picture",
        "Bcrypt password hashing for security",
        "Session management with role-based access control",
        "Account suspension by admin (soft status flag - no DELETE)",
        "Automatic redirect if accessing login/register while already logged in",
    ]),
    ("5.2 Property Listing System", [
        "Full property listings (by landlords): type, address, GPS coordinates, room count, occupancy, gender preference, photos",
        "Peer listings (by students): spare seat/room in their current mess, displayed with 'Student Listed' badge",
        "Mandatory itemized cost form: rent, electricity (individual or shared meter), gas, water, internet, maintenance, caretaker, custom fees",
        "System auto-computes true monthly total from all cost fields",
        "Availability status: Available / Occupied / Soon Vacant (with expected vacate date)",
        "Status badges shown directly on listing cards (green/gold/red)",
        "Occupied listings are hidden from the public search page",
        "Amenity tags: attached bathroom, kitchen, furnished, rooftop, parking, power backup, lift",
        "Photo upload support with multiple images per listing",
    ]),
    ("5.3 Property Search and Discovery", [
        "Filter by zone, rent range, room type (single/shared/full mess/sublet), listing type, amenities",
        "Leaflet.js interactive map with listing pins (UIU campus as default center, Google Maps tiles)",
        "Sort by total monthly cost, composite rating, newest, compatibility score",
        "Personal watchlist - save listings to a shortlist that persists across sessions via database",
        "Real-time filter application without page reload",
    ]),
    ("5.4 Flatmate Compatibility Matching", [
        "Every user fills a preference profile: sleep schedule, study hours, diet, guest policy, smoking tolerance, gender preference, cleanliness (1–5), noise tolerance",
        "Compatibility percentage computed by comparing preferences against current residents using weighted scoring",
        "Compatibility tiers shown on listing cards: High Match (80%+), Good Match (50%+), Low Match (<50%)",
        "Badge color-coded (green/blue/amber) on listing thumbnails",
    ]),
    ("5.5 Seeking Posts Board", [
        "Students post 'Looking For' requests: zone, budget range, room type, gender preference, move-in date, requirements",
        "Other users (landlords, students) can respond to seeking posts",
        "Status: Active or Fulfilled; only active posts shown on main board",
    ]),
    ("5.6 Multi-Dimensional Review System", [
        "Five score dimensions: Value for Money, Accuracy, Landlord Responsiveness, Cleanliness, Safety (each 1–5)",
        "Composite score auto-computed and stored",
        "UNIQUE(listing_id, reviewer_id) - one review per user per listing enforced at database level",
        "Optional text comment (500 chars)",
    ]),
    ("5.7 Complaint and Moderation System", [
        "Categories: Hidden Costs, Harassment, Deposit Not Returned, Listing Misrepresentation, Other",
        "Status workflow: Submitted → Under Review → Resolved (admin-managed)",
        "Optional document/evidence upload with complaint",
        "Admin can view, filter, and resolve complaints from the admin panel",
    ]),
    ("5.8 Landlord Verification System", [
        "Landlords submit NID type, document, and description",
        "Admin approves or rejects → verified badge displayed on all listings",
        "Admin can revoke verification - listings automatically unmarked",
        "Verification status check on Exchange item listing (only verified users can sell)",
    ]),
    ("5.9 Rent History Tracker", [
        "Every rent update archives old value + timestamp into rent_history table (audit trail - never modified)",
        "Listing detail page shows price change timeline as a Chart.js line chart",
    ]),
    ("5.10 Mess Bill Management Module", [
        "Designated mess manager enters monthly utility readings per category",
        "System divides total automatically by number of residents, computing per-person amount",
        "Per-resident payment status tracked (Paid/Unpaid) individually with timestamps",
        "Full monthly bill history accessible from the Dashboard",
    ]),
    ("5.11 UIUNest Exchange - Community Marketplace", [
        "Open to ALL registered users at ANY time for ANY reason",
        "List any household item: furniture, electronics, appliances, kitchen items, study materials",
        "Item fields: category, condition (New/Like New/Good/Fair), price, description, zone, photos",
        "Sold and withdrawn items hidden from public marketplace feed",
        "Structured Offer System: buyer submits formal offer; seller can Accept, Reject, or Counter",
        "All offer states tracked with timestamps: Pending → Countered → Accepted/Rejected/Withdrawn",
        "Atomic transactions (PDO beginTransaction) ensure data consistency",
        "Optional link to a housing listing (items from vacating tenant shown on listing page)",
    ]),
    ("5.12 Admin Dashboard and Analytics", [
        "Manage all users, listings, Exchange items, complaints, verifications from one panel",
        "Analytics: total listings, open complaints, average rent by zone (bar chart), seeking vs. listings demand gap (bar chart)",
        "Approve/reject verification submissions with notification to user",
        "Suspend/reinstate user accounts",
        "Action all complaints with status update",
    ]),
    ("5.13 Notification System", [
        "In-app notifications delivered to specific users (not broadcast to all)",
        "Triggered by: new offer on your item, offer status update, application received, application status update, verification decision",
        "Notification bell with unread count badge in navbar",
        "Mark individual or all notifications as read",
        "No-cache headers prevent browser from serving stale notifications to wrong users",
    ]),
    ("5.14 Application System", [
        "Students apply for available properties with a message to the landlord",
        "Applications tracked in database with status: Pending, Accepted, Rejected",
        "Landlord sees all applications in their dashboard and can update status",
        "Applicant receives notification when status changes",
    ]),
]

for section_title, points in features:
    add_heading(doc, section_title, level=2)
    for point in points:
        p = doc.add_paragraph(f"• {point}")
        p.paragraph_format.left_indent = Cm(0.5)
        if p.runs:
            p.runs[0].font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()

doc.add_page_break()

# ======================== SECTION 6: TECH STACK ========================
add_heading(doc, "6. Technology Stack", level=1)
add_body(doc, "UIUNest is built as a traditional multi-page application (MPA) using the following technologies:")

tech = [
    ("Frontend", "HTML5, Vanilla CSS, Vanilla JavaScript - no frameworks. Custom design system using CSS variables for theming (emerald green/navy color palette). Google Fonts (Inter, Plus Jakarta Sans)."),
    ("Styling", "Custom CSS with CSS custom properties (variables), flexbox, grid, aspect-ratio, animations. No utility frameworks - all styles are hand-written for precision."),
    ("Mapping", "Leaflet.js (v1.9.4) with Google Maps tile layer - interactive property map with custom pins."),
    ("Charts", "Chart.js - used for rent history timeline and admin analytics charts."),
    ("Icons", "Lucide Icons - rendered via JavaScript after page load."),
    ("Backend", "PHP 8+ with PDO (PHP Data Objects) for database interactions. All API endpoints are RESTful PHP files in the /api/ directory."),
    ("Database", "MySQL 8 (via XAMPP locally). PDO with prepared statements throughout - SQL injection prevention enforced."),
    ("Sessions", "PHP native sessions (session_start()) for authentication state. Cache-Control headers added to all session-dependent API endpoints."),
    ("File Uploads", "PHP file upload handling with server-side validation. Files stored in /uploads/ directory with type-specific subdirectories (listings/, exchange/, verification/)."),
    ("Password Security", "PHP password_hash() with PASSWORD_BCRYPT algorithm. password_verify() for authentication."),
    ("Server", "XAMPP (Apache + MySQL + PHP) for local development. Architecture is compatible with Railway or shared hosting deployment."),
]

for label, desc in tech:
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.font.bold = True
    r.font.color.rgb = NAVY
    r.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.size = Pt(11)
    r2.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(6)

add_heading(doc, "File Structure", level=2)
add_code_block(doc, """/UIU-Nest/
├── index.html          - Home page
├── listings.html       - Property search & filter
├── listing-detail.html - Single property detail
├── exchange.html       - Marketplace
├── item-detail.html    - Single item detail
├── seeking.html        - Seeking posts board
├── dashboard.html      - User dashboard (all roles)
├── profile.html        - User profile & preferences
├── admin.html          - Admin panel
├── bills.html          - Mess bill management
├── login.html          - Authentication
├── register.html       - Registration
├── style.css           - Global design system (1600+ lines)
├── app.js              - Shared JS utilities (~900 lines)
├── data.js             - Static mock/fallback data
├── api/
│   ├── db.php          - Database connection (PDO)
│   ├── login.php       - Login endpoint
│   ├── register.php    - Registration endpoint
│   ├── logout.php      - Session destroy
│   ├── me.php          - Current session check
│   ├── listings.php    - CRUD for listings
│   ├── exchange.php    - CRUD for marketplace items
│   ├── offers.php      - Offer state machine
│   ├── applications.php- Housing applications
│   ├── reviews.php     - Review submission
│   ├── seeking.php     - Seeking posts
│   ├── bills.php       - Bill management
│   ├── notifications.php - Per-user notifications
│   ├── watchlist.php   - Save/remove watchlist
│   ├── dashboard.php   - Dashboard data aggregation
│   ├── profile.php     - Profile update
│   ├── upload.php      - File upload handler
│   └── admin_*.php     - Admin operations
├── database/
│   ├── schema.sql      - All 19 CREATE TABLE statements
│   └── seed.sql        - Sample data for testing
└── uploads/            - Uploaded photos""")

doc.add_page_break()

# ======================== SECTION 7: DB DESIGN ========================
add_heading(doc, "7. Database Design Approach", level=1)
add_body(doc, """The UIUNest database is designed around one core principle: every piece of information that a user sees on screen must be stored as a structured, typed, and queryable field in a relational table - not as free text, not as a combined string, and not as an assumption.

The schema follows Third Normal Form (3NF) throughout. Repeating data groups are extracted into their own tables, partial dependencies are eliminated, and transitive dependencies are resolved.

Key design decisions:""")

decisions = [
    ("Single users table", "All three roles (student, landlord, admin) share one table, differentiated by a role ENUM field. This avoids the complexity of separate tables with overlapping structure."),
    ("utility_costs separated from listings", "Utility cost data has a fundamentally different structure and update frequency from listing metadata. A 1-to-1 separation allows the cost form to be updated independently."),
    ("total_monthly stored (not computed on query)", "The most frequently queried field (used in search filtering, sorting, analytics). Stored and updated by the application when any component changes - a deliberate performance trade-off."),
    ("Audit trail pattern for rent_history", "Every UPDATE to base_rent triggers an INSERT into rent_history. Old records are never modified, preserving the complete price change timeline."),
    ("counter_price does not overwrite offer_price", "In the offers table, both buyer's offer and seller's counter are stored in separate columns. The full negotiation trail is preserved for analytics."),
    ("applications table (Table 15)", "A 15th table was added beyond the original design to track student applications to listings, with status lifecycle (pending/accepted/rejected)."),
    ("verifications, watchlists, seeking_responses, notifications added", "Four additional tables were added during implementation (Tables 16–19) to support the verification workflow, saved watchlist functionality, seeking post responses, and the notification system."),
    ("UNIQUE constraints", "UNIQUE(listing_id, reviewer_id) in reviews - one review per user per listing. UNIQUE(user_id, listing_id) in watchlists - no duplicate saves."),
]

for title, text in decisions:
    p = doc.add_paragraph()
    r = p.add_run(f"• {title}: ")
    r.font.bold = True
    r.font.color.rgb = NAVY
    r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)

add_heading(doc, "Entity Groups", level=2)
groups_headers = ["Layer", "Tables", "Purpose"]
groups_rows = [
    ["Identity Layer", "users, zones, user_preferences", "Core entities referenced by all other layers"],
    ["Housing Core", "listings, utility_costs, listing_amenities, rent_history", "Property data with full cost transparency and price audit trail"],
    ["Trust Layer", "reviews, complaints, verifications", "Accountability and verification infrastructure"],
    ["Engagement Layer", "applications, watchlists, seeking_posts, seeking_responses", "User interaction and housing discovery tools"],
    ["Exchange Layer", "items, offers", "Marketplace with full offer state machine"],
    ["Operations Layer", "monthly_bills, bill_payments, notifications", "Mess management and in-app communication"],
]
add_simple_table(doc, groups_headers, groups_rows)
doc.add_page_break()

# ======================== SECTION 8: ERD ========================
add_heading(doc, "8. Entity-Relationship Diagram (ERD)", level=1)
add_body(doc, "The ERD below shows all 19 entities and their relationships. The diagram uses crow's-foot notation: a single bar for 'exactly one' and a crow's foot for 'zero or many.'")
doc.add_paragraph()

erd_path = os.path.join(REPORT_DIR, "New ERD.svg")
alt_path = os.path.join(REPORT_DIR, "new_erd.png")
if os.path.exists(alt_path):
    doc.add_picture(alt_path, width=Inches(6.2))
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("Figure: UIUNest Entity-Relationship Diagram (19 tables)")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.italic = True
else:
    add_body(doc, "[ERD Diagram - insert New ERD.svg or new_erd.png here]", italic=True)

doc.add_page_break()

# ======================== SECTION 9: SCHEMA ========================
add_heading(doc, "9. Database Schema - All 19 Tables", level=1)
add_body(doc, "The following presents the full MySQL CREATE TABLE statements as implemented. All tables use InnoDB with referential integrity enforced via FOREIGN KEY constraints.")

tables_sql = [
    ("Table 1: users", """CREATE TABLE users (
    user_id       INT AUTO_INCREMENT PRIMARY KEY,
    name          VARCHAR(100) NOT NULL,
    email         VARCHAR(150) UNIQUE NOT NULL,
    password_hash VARCHAR(255) NOT NULL,
    role          ENUM('student','landlord','admin') NOT NULL,
    gender        ENUM('male','female','other') NOT NULL,
    phone         VARCHAR(20) NULL,
    university_id VARCHAR(50) NULL,
    profile_pic   VARCHAR(255) NULL,
    status        ENUM('active','suspended') DEFAULT 'active',
    created_at    DATETIME DEFAULT CURRENT_TIMESTAMP
);"""),
    ("Table 2: zones", """CREATE TABLE zones (
    zone_id     INT AUTO_INCREMENT PRIMARY KEY,
    zone_name   VARCHAR(100) NOT NULL,
    description TEXT NULL,
    center_lat  DECIMAL(9,6) NOT NULL,
    center_lng  DECIMAL(9,6) NOT NULL,
    radius_km   DECIMAL(4,2) NOT NULL
);"""),
    ("Table 3: listings", """CREATE TABLE listings (
    listing_id           INT AUTO_INCREMENT PRIMARY KEY,
    user_id              INT,
    zone_id              INT,
    listing_type         ENUM('full_property','peer_listing') NOT NULL,
    property_type        ENUM('single_room','shared_room','full_mess','sublet') NOT NULL,
    title                VARCHAR(200) NOT NULL,
    address              TEXT NOT NULL,
    lat                  DECIMAL(9,6) NOT NULL,
    lng                  DECIMAL(9,6) NOT NULL,
    gender_pref          ENUM('male','female','any') NOT NULL,
    total_rooms          INT NOT NULL,
    current_occupancy    INT DEFAULT 0,
    status               ENUM('available','occupied','soon_vacant') DEFAULT 'available',
    expected_vacate_date DATE NULL,
    is_verified          BOOLEAN DEFAULT FALSE,
    description          TEXT NULL,
    photos               TEXT NULL,
    created_at           DATETIME DEFAULT CURRENT_TIMESTAMP,
    updated_at           DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
    FOREIGN KEY (user_id)  REFERENCES users(user_id) ON DELETE CASCADE,
    FOREIGN KEY (zone_id)  REFERENCES zones(zone_id)
);"""),
    ("Table 4: utility_costs", """CREATE TABLE utility_costs (
    cost_id            INT AUTO_INCREMENT PRIMARY KEY,
    listing_id         INT UNIQUE,
    base_rent          DECIMAL(10,2) NOT NULL,
    electricity_amount DECIMAL(8,2) DEFAULT 0,
    electricity_type   ENUM('individual','shared') NOT NULL,
    gas_bill           DECIMAL(8,2) DEFAULT 0,
    water_bill         DECIMAL(8,2) DEFAULT 0,
    internet_cost      DECIMAL(8,2) DEFAULT 0,
    maintenance_fee    DECIMAL(8,2) DEFAULT 0,
    caretaker_fee      DECIMAL(8,2) DEFAULT 0,
    other_fees         DECIMAL(8,2) DEFAULT 0,
    total_monthly      DECIMAL(10,2) NOT NULL,
    updated_at         DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
    FOREIGN KEY (listing_id) REFERENCES listings(listing_id) ON DELETE CASCADE
);"""),
    ("Table 5: listing_amenities", """CREATE TABLE listing_amenities (
    amenity_id        INT AUTO_INCREMENT PRIMARY KEY,
    listing_id        INT UNIQUE,
    attached_bathroom BOOLEAN DEFAULT FALSE,
    attached_kitchen  BOOLEAN DEFAULT FALSE,
    is_furnished      BOOLEAN DEFAULT FALSE,
    rooftop_access    BOOLEAN DEFAULT FALSE,
    parking           BOOLEAN DEFAULT FALSE,
    power_backup      BOOLEAN DEFAULT FALSE,
    lift_access       BOOLEAN DEFAULT FALSE,
    FOREIGN KEY (listing_id) REFERENCES listings(listing_id) ON DELETE CASCADE
);"""),
    ("Table 6: user_preferences", """CREATE TABLE user_preferences (
    pref_id           INT AUTO_INCREMENT PRIMARY KEY,
    user_id           INT UNIQUE,
    sleep_schedule    ENUM('early','late','flexible') NOT NULL,
    study_hours       INT DEFAULT 0,
    diet              ENUM('vegetarian','non_veg','halal_strict') NOT NULL,
    guest_policy      ENUM('allowed','restricted','not_allowed') NOT NULL,
    smoking_tolerance BOOLEAN DEFAULT FALSE,
    preferred_gender  ENUM('male','female','any') NOT NULL,
    cleanliness_score INT CHECK (cleanliness_score BETWEEN 1 AND 5),
    noise_tolerance   ENUM('quiet','moderate','noisy') NOT NULL,
    FOREIGN KEY (user_id) REFERENCES users(user_id) ON DELETE CASCADE
);"""),
    ("Table 7: reviews", """CREATE TABLE reviews (
    review_id        INT AUTO_INCREMENT PRIMARY KEY,
    listing_id       INT,
    reviewer_id      INT,
    value_for_money  INT CHECK (value_for_money BETWEEN 1 AND 5),
    listing_accuracy INT CHECK (listing_accuracy BETWEEN 1 AND 5),
    landlord_response INT CHECK (landlord_response BETWEEN 1 AND 5),
    cleanliness      INT CHECK (cleanliness BETWEEN 1 AND 5),
    safety           INT CHECK (safety BETWEEN 1 AND 5),
    composite_score  DECIMAL(3,2) NOT NULL,
    comment          VARCHAR(500) NULL,
    created_at       DATETIME DEFAULT CURRENT_TIMESTAMP,
    UNIQUE(listing_id, reviewer_id),
    FOREIGN KEY (listing_id)   REFERENCES listings(listing_id) ON DELETE CASCADE,
    FOREIGN KEY (reviewer_id)  REFERENCES users(user_id) ON DELETE CASCADE
);"""),
    ("Table 8: complaints", """CREATE TABLE complaints (
    complaint_id    INT AUTO_INCREMENT PRIMARY KEY,
    complainant_id  INT,
    against_user_id INT,
    listing_id      INT NULL,
    category        ENUM('hidden_costs','harassment',
                         'deposit_not_returned','misrepresentation','other') NOT NULL,
    description     TEXT NOT NULL,
    status          ENUM('submitted','under_review','resolved') DEFAULT 'submitted',
    created_at      DATETIME DEFAULT CURRENT_TIMESTAMP,
    resolved_at     DATETIME NULL,
    FOREIGN KEY (complainant_id)  REFERENCES users(user_id) ON DELETE CASCADE,
    FOREIGN KEY (against_user_id) REFERENCES users(user_id) ON DELETE CASCADE,
    FOREIGN KEY (listing_id)      REFERENCES listings(listing_id) ON DELETE CASCADE
);"""),
    ("Table 9: seeking_posts", """CREATE TABLE seeking_posts (
    post_id          INT AUTO_INCREMENT PRIMARY KEY,
    user_id          INT,
    zone_id          INT,
    budget_min       DECIMAL(10,2) NOT NULL,
    budget_max       DECIMAL(10,2) NOT NULL,
    property_type    ENUM('single_room','shared_room','full_mess','sublet','any') NOT NULL,
    preferred_gender ENUM('male','female','any') NOT NULL,
    move_in_date     DATE NULL,
    requirements     TEXT NULL,
    status           ENUM('active','fulfilled') DEFAULT 'active',
    created_at       DATETIME DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (user_id)  REFERENCES users(user_id) ON DELETE CASCADE,
    FOREIGN KEY (zone_id)  REFERENCES zones(zone_id)
);"""),
    ("Table 10: monthly_bills", """CREATE TABLE monthly_bills (
    bill_id            INT AUTO_INCREMENT PRIMARY KEY,
    listing_id         INT,
    bill_month         DATE NOT NULL,
    electricity_amount DECIMAL(8,2) DEFAULT 0,
    gas_amount         DECIMAL(8,2) DEFAULT 0,
    water_amount       DECIMAL(8,2) DEFAULT 0,
    internet_amount    DECIMAL(8,2) DEFAULT 0,
    other_amount       DECIMAL(8,2) DEFAULT 0,
    total_amount       DECIMAL(10,2) NOT NULL,
    per_person_amount  DECIMAL(10,2) NOT NULL,
    created_by         INT,
    created_at         DATETIME DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (listing_id) REFERENCES listings(listing_id) ON DELETE CASCADE,
    FOREIGN KEY (created_by) REFERENCES users(user_id) ON DELETE CASCADE
);"""),
    ("Table 11: bill_payments", """CREATE TABLE bill_payments (
    payment_id       INT AUTO_INCREMENT PRIMARY KEY,
    bill_id          INT,
    resident_user_id INT NULL,
    resident_label   VARCHAR(50) NULL,
    status           ENUM('paid','unpaid') DEFAULT 'unpaid',
    paid_at          DATETIME NULL,
    FOREIGN KEY (bill_id)          REFERENCES monthly_bills(bill_id) ON DELETE CASCADE,
    FOREIGN KEY (resident_user_id) REFERENCES users(user_id) ON DELETE CASCADE
);"""),
    ("Table 12: rent_history", """CREATE TABLE rent_history (
    history_id INT AUTO_INCREMENT PRIMARY KEY,
    listing_id INT,
    old_rent   DECIMAL(10,2) NOT NULL,
    new_rent   DECIMAL(10,2) NOT NULL,
    changed_at DATETIME DEFAULT CURRENT_TIMESTAMP,
    changed_by INT,
    FOREIGN KEY (listing_id)  REFERENCES listings(listing_id) ON DELETE CASCADE,
    FOREIGN KEY (changed_by)  REFERENCES users(user_id) ON DELETE CASCADE
);"""),
    ("Table 13: items (UIUNest Exchange)", """CREATE TABLE items (
    item_id            INT AUTO_INCREMENT PRIMARY KEY,
    seller_id          INT,
    zone_id            INT,
    listing_id         INT NULL,
    category           ENUM('furniture','appliances','electronics',
                            'kitchen','study','other') NOT NULL,
    title              VARCHAR(200) NOT NULL,
    description        TEXT NULL,
    item_condition     ENUM('new','like_new','good','fair') NOT NULL,
    asking_price       DECIMAL(10,2) NOT NULL,
    reason_for_selling VARCHAR(300) NULL,
    photo_url          VARCHAR(255) NULL,
    status             ENUM('available','sold','withdrawn') DEFAULT 'available',
    created_at         DATETIME DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (seller_id) REFERENCES users(user_id) ON DELETE CASCADE,
    FOREIGN KEY (zone_id)   REFERENCES zones(zone_id),
    FOREIGN KEY (listing_id) REFERENCES listings(listing_id) ON DELETE SET NULL
);"""),
    ("Table 14: offers", """CREATE TABLE offers (
    offer_id      INT AUTO_INCREMENT PRIMARY KEY,
    item_id       INT,
    buyer_id      INT,
    offer_price   DECIMAL(10,2) NOT NULL,
    message       VARCHAR(300) NULL,
    status        ENUM('pending','countered','accepted','rejected','withdrawn') DEFAULT 'pending',
    counter_price DECIMAL(10,2) NULL,
    created_at    DATETIME DEFAULT CURRENT_TIMESTAMP,
    updated_at    DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
    FOREIGN KEY (item_id)   REFERENCES items(item_id) ON DELETE CASCADE,
    FOREIGN KEY (buyer_id)  REFERENCES users(user_id) ON DELETE CASCADE
);"""),
    ("Table 15: applications", """CREATE TABLE applications (
    application_id INT AUTO_INCREMENT PRIMARY KEY,
    listing_id     INT,
    applicant_id   INT,
    message        TEXT NULL,
    status         ENUM('pending','accepted','rejected') DEFAULT 'pending',
    created_at     DATETIME DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (listing_id)    REFERENCES listings(listing_id) ON DELETE CASCADE,
    FOREIGN KEY (applicant_id)  REFERENCES users(user_id) ON DELETE CASCADE
);"""),
    ("Table 16: verifications", """CREATE TABLE verifications (
    verification_id INT AUTO_INCREMENT PRIMARY KEY,
    user_id         INT,
    nid_type        ENUM('National ID','Passport','Driving License') NOT NULL,
    document_path   VARCHAR(255) NOT NULL,
    description     TEXT NULL,
    status          ENUM('pending','approved','rejected') DEFAULT 'pending',
    submitted_at    DATETIME DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (user_id) REFERENCES users(user_id) ON DELETE CASCADE
);"""),
    ("Table 17: watchlists", """CREATE TABLE watchlists (
    watchlist_id INT AUTO_INCREMENT PRIMARY KEY,
    user_id      INT,
    listing_id   INT,
    added_at     DATETIME DEFAULT CURRENT_TIMESTAMP,
    UNIQUE(user_id, listing_id),
    FOREIGN KEY (user_id)    REFERENCES users(user_id) ON DELETE CASCADE,
    FOREIGN KEY (listing_id) REFERENCES listings(listing_id) ON DELETE CASCADE
);"""),
    ("Table 18: seeking_responses", """CREATE TABLE seeking_responses (
    response_id  INT AUTO_INCREMENT PRIMARY KEY,
    post_id      INT,
    responder_id INT,
    message      TEXT NULL,
    status       ENUM('pending','accepted','rejected') DEFAULT 'pending',
    created_at   DATETIME DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (post_id)      REFERENCES seeking_posts(post_id) ON DELETE CASCADE,
    FOREIGN KEY (responder_id) REFERENCES users(user_id) ON DELETE CASCADE
);"""),
    ("Table 19: notifications", """CREATE TABLE notifications (
    notif_id   INT AUTO_INCREMENT PRIMARY KEY,
    user_id    INT,
    type       VARCHAR(50),
    message    TEXT,
    is_read    BOOLEAN DEFAULT FALSE,
    link       VARCHAR(255) NULL,
    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (user_id) REFERENCES users(user_id) ON DELETE CASCADE
);"""),
]

for title, sql in tables_sql:
    add_heading(doc, title, level=2)
    add_code_block(doc, sql)

doc.add_page_break()

# ======================== SECTION 10: SQL QUERIES ========================
add_heading(doc, "10. Key SQL Queries", level=1)
add_body(doc, "This section presents the core SQL queries used by UIUNest, organized by DBMS operation type.")

add_heading(doc, "10.1 SELECT with Multi-Table JOIN", level=2)
add_body(doc, "Fetching all listing data in one query joining 4 tables:")
add_code_block(doc, """-- api/listings.php: Full listings query with owner, zone, and preferences
SELECT l.*, 
       z.zone_name AS zone, 
       u.name AS owner_name, 
       u.email AS owner_email,
       up.sleep_schedule AS sleep, 
       up.diet, 
       up.guest_policy AS guest,
       up.smoking_tolerance AS smoking, 
       up.noise_tolerance AS noise, 
       up.cleanliness_score AS cleanliness
FROM listings l
LEFT JOIN zones z           ON l.zone_id = z.zone_id
LEFT JOIN users u           ON l.user_id = u.user_id
LEFT JOIN user_preferences up ON l.user_id = up.user_id
WHERE l.listing_id = ?;""")

add_body(doc, "Fetching reviews joined with reviewer names:")
add_code_block(doc, """-- api/listings.php: Reviews with reviewer name
SELECT r.*, u.name AS reviewer_name 
FROM reviews r
LEFT JOIN users u ON r.reviewer_id = u.user_id
ORDER BY r.created_at DESC;""")

add_body(doc, "Exchange items with seller name and zone:")
add_code_block(doc, """-- api/exchange.php: Items query
SELECT i.*, u.name AS seller, u.email AS seller_email, z.zone_name AS zone 
FROM items i
LEFT JOIN users u ON i.seller_id = u.user_id
LEFT JOIN zones z ON i.zone_id = z.zone_id
ORDER BY i.created_at DESC;""")

add_heading(doc, "10.2 SELECT with Aggregate Functions", level=2)
add_body(doc, "Average rent per zone (used on Admin analytics dashboard):")
add_code_block(doc, """-- api/admin_stats.php: Average rent by zone
SELECT z.zone_name AS zone, 
       COALESCE(AVG(uc.total_monthly), 0) AS avg_rent
FROM zones z
LEFT JOIN listings l       ON z.zone_id = l.zone_id
LEFT JOIN utility_costs uc ON l.listing_id = uc.listing_id
GROUP BY z.zone_id, z.zone_name;""")

add_body(doc, "Composite review score calculation:")
add_code_block(doc, """-- api/reviews.php: Compute composite score
SELECT listing_id,
       AVG((value_for_money + listing_accuracy + landlord_response + 
            cleanliness + safety) / 5.0) AS composite_score,
       COUNT(*) AS total_reviews
FROM reviews
WHERE listing_id = ?
GROUP BY listing_id;""")

add_heading(doc, "10.3 SELECT with Subquery", level=2)
add_body(doc, "Seeking vs listings demand gap per zone (correlated subquery):")
add_code_block(doc, """-- api/admin_stats.php: Demand vs supply gap
SELECT 
    z.zone_name AS zone,
    (SELECT COUNT(*) FROM seeking_posts sp 
     WHERE sp.zone_id = z.zone_id AND sp.status = 'active') AS seeking,
    (SELECT COUNT(*) FROM listings l 
     WHERE l.zone_id = z.zone_id) AS listings
FROM zones z;""")

add_body(doc, "Check for duplicate active offer before INSERT:")
add_code_block(doc, """-- api/offers.php: Prevent duplicate offer submission
SELECT offer_id FROM offers 
WHERE item_id = ? 
  AND buyer_id = ? 
  AND status IN ('pending', 'countered', 'accepted');""")

add_heading(doc, "10.4 INSERT Operations", level=2)
add_body(doc, "User registration with bcrypt password:")
add_code_block(doc, """-- api/register.php: Insert new user
INSERT INTO users (name, email, password_hash, role, gender, university_id)
VALUES (?, ?, ?, ?, ?, ?);""")

add_body(doc, "Listing creation with cost breakdown (two separate INSERTs in sequence):")
add_code_block(doc, """-- api/listings.php: Create listing
INSERT INTO listings 
  (user_id, zone_id, listing_type, property_type, title, address, 
   lat, lng, gender_pref, total_rooms, description, photos)
VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?);

-- Then insert associated costs
INSERT INTO utility_costs 
  (listing_id, base_rent, electricity_amount, electricity_type,
   gas_bill, water_bill, internet_cost, maintenance_fee, 
   caretaker_fee, other_fees, total_monthly)
VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?);""")

add_body(doc, "Offer submission using atomic transaction:")
add_code_block(doc, """-- api/offers.php: Insert offer with notification (atomic)
BEGIN TRANSACTION;

INSERT INTO offers (item_id, buyer_id, offer_price, message)
VALUES (?, ?, ?, ?);

INSERT INTO notifications (user_id, type, message, link)
VALUES (?, 'offer', ?, 'dashboard.html?tab=offers');

COMMIT;""")

add_heading(doc, "10.5 UPDATE Operations", level=2)
add_body(doc, "Updating listing status (Available / Occupied / Soon Vacant):")
add_code_block(doc, """-- api/update_status.php: Change listing availability
UPDATE listings 
SET status = ?, expected_vacate_date = ?, updated_at = NOW()
WHERE listing_id = ? AND user_id = ?;""")

add_body(doc, "Offer state machine - seller responds to an offer:")
add_code_block(doc, """-- api/offers.php: Seller accepts, rejects, or counters
UPDATE offers 
SET status = ?, counter_price = ?, updated_at = NOW()
WHERE offer_id = ?;

-- If accepted: mark item as sold
UPDATE items SET status = 'sold'
WHERE item_id = ?;""")

add_body(doc, "Rent history audit trail - archive old rent before updating:")
add_code_block(doc, """-- api/listings.php: Rent update with history
INSERT INTO rent_history (listing_id, old_rent, new_rent, changed_by)
SELECT listing_id, base_rent, ?, ?
FROM utility_costs WHERE listing_id = ?;

UPDATE utility_costs
SET base_rent = ?, total_monthly = ?, updated_at = NOW()
WHERE listing_id = ?;""")

add_heading(doc, "10.6 DELETE Operations", level=2)
add_body(doc, "Admin suspending a user account (soft delete - status flag, not physical DELETE):")
add_code_block(doc, """-- api/admin_action_user.php: Suspend user (soft delete)
UPDATE users SET status = 'suspended' WHERE user_id = ?;""")

add_body(doc, "Hard delete for exchange items (cascade deletes associated offers):")
add_code_block(doc, """-- api/exchange.php: Delete item (ON DELETE CASCADE handles offers)
DELETE FROM items WHERE item_id = ? AND seller_id = ?;""")

add_body(doc, "Remove from watchlist:")
add_code_block(doc, """-- api/watchlist.php: Toggle watchlist (DELETE if exists)
DELETE FROM watchlists WHERE user_id = ? AND listing_id = ?;""")

add_heading(doc, "10.7 Notifications Query (User-Specific)", level=2)
add_body(doc, "Fetching notifications for the currently logged-in user only:")
add_code_block(doc, """-- api/notifications.php: Fetch user-specific notifications
SELECT * FROM notifications 
WHERE user_id = ? 
ORDER BY created_at DESC 
LIMIT 50;

-- Mark all as read
UPDATE notifications SET is_read = TRUE WHERE user_id = ?;""")

doc.add_page_break()

# ======================== SECTION 11: SCREENSHOTS ========================
add_heading(doc, "11. Application Screenshots", level=1)
add_body(doc, "The following screenshots show the current live version of UIUNest running locally on XAMPP.")

screenshots = [
    ("screenshot_home.png", "UIUNest Home Page - Hero section with navigation bar, stats, and featured zones", 5.8),
    ("screenshot_login.png", "Login Page - Clean authentication form with emerald brand theme", 4.5),
    ("screenshot_register.png", "Registration Page - Role selection (Student/Landlord) with university ID field", 4.5),
    ("screenshot_listings.png", "Listings Page - Filter sidebar, property cards with status badges, and interactive map", 5.8),
    ("screenshot_exchange.png", "UIUNest Exchange Marketplace - Product grid with category and condition filters", 5.8),
    ("screenshot_seeking.png", "Seeking Board - Student housing request cards with zone and budget information", 5.8),
]

for filename, caption, width in screenshots:
    add_screenshot(doc, filename, caption, width)

doc.add_page_break()

# ======================== SECTION 12: USER ROLES ========================
add_heading(doc, "12. User Roles & Access Control", level=1)
add_body(doc, "UIUNest implements role-based access control (RBAC) through a single users table differentiated by a role ENUM field. Each role has distinct permissions enforced at both the UI and API layer.")

roles_headers = ["Feature / Action", "Student", "Landlord", "Admin"]
roles_rows = [
    ["Register & login", "✓", "✓", "✓"],
    ["View listings", "✓", "✓", "✓"],
    ["Apply for a listing", "✓", "✗", "✗"],
    ["Create full property listing", "✗", "✓", "✗"],
    ["Create peer listing (spare seat)", "✓", "✗", "✗"],
    ["Submit verification document", "✗", "✓", "✗"],
    ["Sell item on Exchange", "✓ (if verified)", "✓ (if verified)", "✗"],
    ["Submit offer on Exchange", "✓", "✓", "✗"],
    ["Submit a review", "✓", "✗", "✗"],
    ["File a complaint", "✓", "✓", "✗"],
    ["Manage mess bills", "✓ (own mess)", "✓ (own mess)", "✗"],
    ["Approve verifications", "✗", "✗", "✓"],
    ["Resolve complaints", "✗", "✗", "✓"],
    ["Suspend users", "✗", "✗", "✓"],
    ["View admin analytics", "✗", "✗", "✓"],
    ["Access admin panel", "✗", "✗", "✓"],
]
add_simple_table(doc, roles_headers, roles_rows)
doc.add_page_break()

# ======================== SECTION 13: ARCHITECTURE ========================
add_heading(doc, "13. System Architecture", level=1)
add_body(doc, "UIUNest follows a classic three-tier web application architecture:")

arch_text = [
    ("Presentation Layer (Frontend)", "HTML/CSS/JS pages served as static files by Apache. No server-side rendering - all dynamic content is loaded via JavaScript fetch() calls to API endpoints. The design system is defined in style.css using CSS custom properties. Shared utilities (authentication helpers, rendering functions, toast notifications, watchlist management) are in app.js which is included on every page."),
    ("Application Layer (PHP API)", "RESTful API endpoints in the /api/ directory. Each file handles one resource (listings.php, exchange.php, offers.php, etc.). All endpoints use PDO prepared statements for SQL injection prevention. Session state (session_start()) identifies the current user. Role checks are enforced at the API level - a student cannot call admin endpoints even by URL manipulation."),
    ("Data Layer (MySQL)", "MySQL 8 database with 19 tables organized into functional layers. Foreign key constraints with ON DELETE CASCADE ensure referential integrity. UNIQUE constraints prevent duplicate reviews, duplicate watchlist entries, and duplicate active offers. All monetary values use DECIMAL(10,2) - never FLOAT - to avoid rounding errors."),
]

for title, text in arch_text:
    add_heading(doc, title, level=2)
    add_body(doc, text)

add_heading(doc, "Request Flow Example: Submitting an Offer", level=2)
add_code_block(doc, """Browser (JavaScript)
  └─► fetch('api/offers.php', { method: 'POST', body: JSON.stringify({itemId, offerPrice, message}) })
        │
        ▼
PHP (api/offers.php)
  ├─► session_start() → validate user is logged in
  ├─► Check for existing active offer (SELECT query)
  ├─► BEGIN TRANSACTION
  │     ├─► INSERT INTO offers (item_id, buyer_id, offer_price, message)
  │     └─► INSERT INTO notifications (user_id=seller, type='offer', message, link)
  └─► COMMIT → return { success: true }
        │
        ▼
Browser
  └─► Show success toast → update UI state""")

doc.add_page_break()

# ======================== SECTION 14: LIMITATIONS ========================
add_heading(doc, "14. Limitations", level=1)
limitations = [
    ("No real-time messaging", "The current platform uses a request-response model. There is no WebSocket or real-time push system - users must refresh to see new notifications and offers. A future version could use Server-Sent Events or WebSockets."),
    ("File storage is local", "Uploaded photos are stored on the local server filesystem (/uploads/). In a production deployment, these would need to be moved to a cloud storage service (e.g., AWS S3, Cloudflare R2) to support horizontal scaling."),
    ("Compatibility matching is client-side", "The compatibility score is computed in JavaScript on the frontend (app.js) rather than in an SQL query. This means it cannot be used as a sort criterion for server-side queries. A production version would implement this as a stored function or computed column."),
    ("No email/SMS integration", "Notifications exist only inside the platform. There is no email or SMS system for critical alerts (new offer, application accepted). Users must log in to see notifications."),
    ("Map is for display only", "The Leaflet map shows listing pins but does not allow drawing a search radius or filtering by map viewport. Zone-based filtering is the primary geographic filter."),
    ("No payment gateway", "UIUNest tracks bill payment status (paid/unpaid) but has no integration with a payment gateway (bKash, Nagad, or card). All payments are assumed to be recorded manually by the mess manager."),
    ("Verification uses document metadata only", "The prototype stores document type and description for verification. A production system would require actual ID document upload and an admin review workflow with document image viewing."),
]
for title, text in limitations:
    p = doc.add_paragraph()
    r = p.add_run(f"• {title}: ")
    r.font.bold = True
    r.font.color.rgb = NAVY
    r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(8)

doc.add_page_break()

# ======================== SECTION 15: FUTURE WORK ========================
add_heading(doc, "15. Future Work", level=1)
futures = [
    ("Mobile Application", "A native Android/iOS app using the existing PHP API as backend, providing push notifications for offers, applications, and availability changes."),
    ("Real-Time Notifications", "WebSocket integration (e.g., Ratchet PHP or Node.js bridge) for instant offer and application notifications without page refresh."),
    ("Advanced Compatibility Matching", "Move the compatibility scoring algorithm from JavaScript to a MySQL stored procedure, enabling server-side sorting and filtering by compatibility score."),
    ("Expansion to Other Universities", "The zone system is already designed for multi-campus support. Adding more universities requires only a new set of zones and allowing university-specific domain email registration."),
    ("AI-Powered Listing Recommendations", "A recommendation engine using collaborative filtering (users who liked X also viewed Y) built on top of the existing watchlist and view history data."),
    ("Payment Integration", "Integration with bKash or Nagad APIs to allow mess members to pay their monthly bill share directly through the platform, with automatic status update in bill_payments."),
    ("Neighborhood Rating System", "The neighborhood_ratings table is designed in the schema but not yet implemented. Adding the ability for users to rate zones on safety, transport, and noise would add another data layer to the platform."),
    ("Moderation AI", "Using a content moderation API (e.g., Google Perspective API) to automatically flag offensive content in complaints and review text before admin review."),
]
for title, text in futures:
    p = doc.add_paragraph()
    r = p.add_run(f"• {title}: ")
    r.font.bold = True
    r.font.color.rgb = NAVY
    r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(8)

doc.add_page_break()

# ======================== SECTION 16: CONCLUSION ========================
add_heading(doc, "16. Conclusion", level=1)
add_body(doc, """UIUNest was built to solve a real problem that every UIU student living near campus has faced: finding housing through an informal, fragmented, and often dishonest process. The platform replaces that process with a structured, transparent, and accountable system.

From a database systems perspective, UIUNest demonstrates the full breadth of relational database concepts covered in the DBMS laboratory course:

  • SELECT with multi-table LEFT JOINs across up to 4 tables in a single query
  • INSERT operations including atomic multi-table inserts using PDO transactions
  • UPDATE operations including the audit trail pattern (rent_history) and state machine transitions (offer status)
  • DELETE with ON DELETE CASCADE for referential integrity
  • Aggregate functions: AVG(), COUNT(), SUM(), COALESCE() for analytics
  • Correlated subqueries for demand gap analysis
  • UNIQUE constraints enforced at the database level for business rules
  • Foreign key constraints across all 19 tables with appropriate cascade/set-null behaviors

The system was built with a production-grade mindset: all SQL uses prepared statements, passwords are bcrypt-hashed, sessions are validated at every API endpoint, and critical operations (offer submission, bill entry) use atomic transactions to prevent partial state.

UIUNest demonstrates that a DBMS-driven application can solve real, everyday problems for a university community - while simultaneously serving as a comprehensive exercise in relational database design, normalization, and SQL query engineering.""")

# ======================== SAVE ========================
doc.save(OUTPUT_PATH)
print(f"Report saved to: {OUTPUT_PATH}")
